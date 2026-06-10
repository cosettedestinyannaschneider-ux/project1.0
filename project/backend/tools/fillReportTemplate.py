import json
import os
import re
import sys

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt
from docx.text.paragraph import Paragraph


LABELS = {
    "compiler_unit": "\u7f16\u5236\u5355\u4f4d",
    "compiler": "\u7f16\u5236\u4eba\u5458",
    "auditor": "\u5ba1\u6838\u4eba\u5458",
    "date": "\u7f16\u5236\u65e5\u671f",
    "phone": "\u8054\u7cfb\u7535\u8bdd",
    "email": "\u7535\u5b50\u90ae\u7bb1",
    "placeholder_company": "XXXX\u96c6\u56e2\u80a1\u4efd\u6709\u9650\u516c\u53f8",
    "placeholder_project": "XXX\u9879\u76ee\u90e8",
    "sample_project": "\u9526\u5510\u5e9c\u9879\u76ee\u90e8",
    "sample_project_reversed": "\u90e8\u9526\u5510\u5e9c\u9879\u76ee",
    "opinion_title": "\u4e8c\u3001\u9690\u60a3\u6392\u67e5\u7efc\u5408\u610f\u89c1",
    "opinion_sub_1": "\uff08\u4e00\uff09 \u5b89\u5168\u7ba1\u7406\u6539\u8fdb\u65b9\u5411\u53ca\u5efa\u8bae",
    "opinion_sub_2": "\uff08\u4e8c\uff09 \u7efc\u5408\u5efa\u8bae",
    "hazard_table_header": [
        "\u4f01\u4e1a\n\u5730\u70b9",
        "\u73b0\u573a\u5b58\u5728\u95ee\u9898\uff08\u9690\u60a3\u63cf\u8ff0\u53ca\u56fe\u7247\uff09",
        "\u9690\u60a3\n\u7b49\u7ea7",
        "\u6574\u6539\u5efa\u8bae",
        "\u8d23\u4efb\n\u5212\u5206",
    ],
}


def clean(value):
    return str(value or "").replace("\u25a1", "").replace("\u25a0", "").strip()


def normalize_standard_code(value):
    value = re.sub(r"\bGB\s*/\s*T\s*[- ]?\s*", "GB/T ", value, flags=re.I)
    value = re.sub(r"\bDB(\d+)\s*/\s*T\s*[- ]?\s*", r"DB\1/T ", value, flags=re.I)
    value = re.sub(r"\b(JGJ|AQ|TSG|XF|GA|JG|CJ|SY|GB)\s*-\s*", r"\1 ", value, flags=re.I)
    value = re.sub(r"\s+", " ", value)
    return value.strip()


def normalize_reference_object(value):
    if not isinstance(value, dict):
        return ""
    name = clean(value.get("name") or value.get("title") or value.get("standard_name"))
    code = normalize_standard_code(clean(value.get("code") or value.get("standard_code") or value.get("number")))
    clause = clean(value.get("clause") or value.get("article") or value.get("item"))
    content = clean(value.get("content") or value.get("clause_content") or value.get("text"))
    if not name and not code and not clause and not content:
        return ""
    prefix = f"《{name}》" if name and not name.startswith("《") else name
    if code and code not in prefix:
        prefix = f"{prefix}（{code}）" if prefix else f"（{code}）"
    if clause and not clause.startswith(("第", "条")):
        clause = f"第{clause}条" if re.match(r"^\d", clause) else clause
    parts = [part for part in [prefix, clause] if part]
    reference = "".join(parts)
    if content:
        reference = f"{reference}：{content}" if reference else content
    return reference


def normalize_reference_text(value):
    if isinstance(value, dict):
        value = normalize_reference_object(value)
    value = normalize_standard_code(clean(value))
    value = re.sub(r"^\s*\d+\s*[.．、]\s*(?:[→>]\s*)?", "", value)
    value = value.strip(" ；;。")
    if not value:
        return ""

    value = value.replace("《《", "《").replace("》》", "》")
    value = re.sub(r"（\s*）", "", value)

    if "《" not in value:
        code_first = re.match(
            r"^((?:GB/T|GB|JGJ|AQ|DB\d*/T|TSG|XF|GA|JG|CJ|SY)\s*[A-Z0-9./-]+-\d{4})\s+(.+)$",
            value,
            flags=re.I,
        )
        if code_first:
            code = normalize_standard_code(code_first.group(1))
            title_and_clause = clean(code_first.group(2)).strip(" ；;。")
            clause_match = re.match(
                r"^(.+?)\s*((?:第[\d一二三四五六七八九十百千万零〇两]+[条款项节]|第?[\d.]+[条款项节])[:：].+)$",
                title_and_clause,
            )
            if clause_match:
                title = clause_match.group(1).strip(" ；;。")
                clause = clause_match.group(2).strip(" ；;。")
                value = f"《{title}》（{code}）{clause}"
            else:
                title = title_and_clause
                value = f"《{title}》（{code}）"
        else:
            title_code = re.match(
                r"^(.+?)\s*[（(]\s*((?:GB/T|GB|JGJ|AQ|DB\d*/T|TSG|XF|GA|JG|CJ|SY)\s*[A-Z0-9./-]+-\d{4})\s*[）)](.*)$",
                value,
                flags=re.I,
            )
            if title_code:
                title = clean(title_code.group(1)).strip(" ；;。")
                code = normalize_standard_code(title_code.group(2))
                rest = clean(title_code.group(3)).strip(" ；;。")
                value = f"《{title}》（{code}）{rest}"
            elif any(token in value for token in ["法", "条例", "规范", "标准", "规程", "导则", "规定", "办法"]):
                value = f"《{value}》"
    else:
        value = re.sub(
            r"《([^》]+)》\s*[（(]\s*((?:GB/T|GB|JGJ|AQ|DB\d*/T|TSG|XF|GA|JG|CJ|SY)\s*[A-Z0-9./-]+-\d{4})\s*[）)]",
            lambda match: f"《{match.group(1)}》（{normalize_standard_code(match.group(2))}）",
            value,
            flags=re.I,
        )

    value = re.sub(r"\s*[:：]\s*", "：", value)
    value = re.sub(r"\s+", " ", value)

    return f"{value}；"


def remove_element(element):
    parent = element.getparent()
    if parent is not None:
        parent.remove(element)


def remove_paragraph(paragraph):
    remove_element(paragraph._element)


def copy_paragraph_format(source, target):
    target.style = source.style
    target.alignment = source.alignment
    src = source.paragraph_format
    dst = target.paragraph_format
    dst.left_indent = src.left_indent
    dst.right_indent = src.right_indent
    dst.first_line_indent = src.first_line_indent
    dst.space_before = src.space_before
    dst.space_after = src.space_after
    dst.line_spacing = src.line_spacing
    dst.line_spacing_rule = src.line_spacing_rule


def copy_run_format(source, target):
    if source is None:
        target.font.name = "仿宋"
        target._element.rPr.rFonts.set(qn("w:eastAsia"), "仿宋")
        target.font.size = Pt(15)
        return
    target.bold = source.bold
    target.italic = source.italic
    target.underline = source.underline
    target.font.name = source.font.name
    target.font.size = source.font.size
    if source.font.name:
        target._element.rPr.rFonts.set(qn("w:eastAsia"), source.font.name)


def insert_paragraph_after(paragraph, value="", first_line_indent_cm=None, template_paragraph=None):
    new_element = OxmlElement("w:p")
    paragraph._p.addnext(new_element)
    new_paragraph = Paragraph(new_element, paragraph._parent)
    if template_paragraph is not None:
        copy_paragraph_format(template_paragraph, new_paragraph)
    if first_line_indent_cm is not None:
        new_paragraph.paragraph_format.first_line_indent = Cm(first_line_indent_cm)
    if value:
        run = new_paragraph.add_run(value)
        source_run = template_paragraph.runs[0] if template_paragraph is not None and template_paragraph.runs else None
        copy_run_format(source_run, run)
    return new_paragraph


def clear_cell(cell):
    for paragraph in list(cell.paragraphs):
        remove_paragraph(paragraph)
    cell.add_paragraph()


def set_cell_width(cell, width_cm):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(width_cm * 567)))
    tc_w.set(qn("w:type"), "dxa")


def set_cell_text(cell, value, bold=False, align=None, font_size=10.5):
    clear_cell(cell)
    paragraph = cell.paragraphs[0]
    if align is not None:
        paragraph.alignment = align
    run = paragraph.add_run(clean(value))
    run.bold = bold
    run.font.size = Pt(font_size)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def replace_runs(paragraph, replacements):
    for run in paragraph.runs:
        value = run.text
        for old, new in replacements.items():
            value = value.replace(old, new)
        run.text = value


def replace_all(document, replacements):
    for paragraph in document.paragraphs:
        replace_runs(paragraph, replacements)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    replace_runs(paragraph, replacements)


def fill_cover_table(document, payload):
    if not document.tables:
        return
    enterprise = payload.get("enterprise", {})
    values = {
        LABELS["compiler_unit"]: payload.get("compilerUnit") or "XXXX\u5b89\u5168\u6280\u672f\u54a8\u8be2\u6709\u9650\u516c\u53f8",
        LABELS["compiler"]: enterprise.get("inspector_name") or payload.get("auditorName") or "",
        LABELS["auditor"]: payload.get("auditorName") or "",
        LABELS["date"]: payload.get("inspectionDate") or "",
        LABELS["phone"]: enterprise.get("phone") or "",
        LABELS["email"]: enterprise.get("email") or "",
    }
    table = document.tables[0]
    for row in table.rows:
        if len(row.cells) < 2:
            continue
        label = clean(row.cells[0].text).replace(":", "").replace("\uff1a", "")
        for key, value in values.items():
            if key in label:
                set_cell_text(row.cells[1], value)


def item_image_path(item, index, image_paths):
    try:
        image_index = int(item.get("image_id") or 0) - 1
    except Exception:
        image_index = -1
    if 0 <= image_index < len(image_paths):
        return image_paths[image_index]
    if index < len(image_paths):
        return image_paths[index]
    return None


def build_clean_hazard_table(document, payload):
    old_table = document.tables[1] if len(document.tables) > 1 else None
    table = document.add_table(rows=1, cols=5)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.style = "Table Grid"

    widths = [1.7, 7.3, 1.55, 5.9, 1.65]
    headers = LABELS["hazard_table_header"]
    for index, cell in enumerate(table.rows[0].cells):
        set_cell_width(cell, widths[index])
        set_cell_text(cell, headers[index], bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, font_size=10)

    enterprise = payload.get("enterprise", {})
    items = payload.get("items") or []
    image_paths = payload.get("imagePaths") or []
    project_name = enterprise.get("project_name") or enterprise.get("region") or "\u73b0\u573a"

    for index, item in enumerate(items):
        row = table.add_row()
        for cell_index, cell in enumerate(row.cells):
            set_cell_width(cell, widths[cell_index])

        set_cell_text(row.cells[0], item.get("location") or project_name, font_size=9.5)

        evidence_cell = row.cells[1]
        clear_cell(evidence_cell)
        evidence_cell.paragraphs[0].add_run(item.get("hazard_description") or "\u5f85\u8865\u5145")
        evidence_cell.paragraphs[0].runs[0].font.size = Pt(9.5)
        image_path = item_image_path(item, index, image_paths)
        if image_path and os.path.exists(image_path):
            paragraph = evidence_cell.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = paragraph.add_run()
            run.add_picture(image_path, width=Inches(2.0))

        set_cell_text(row.cells[2], item.get("hazard_level") or "\u4e00\u822c\u9690\u60a3", align=WD_ALIGN_PARAGRAPH.CENTER, font_size=9.5)
        set_cell_text(row.cells[3], item.get("suggestion") or "\u5efa\u8bae\u7ed3\u5408\u73b0\u573a\u60c5\u51b5\u8865\u5145\u5177\u4f53\u6574\u6539\u63aa\u65bd\u3002", font_size=9.5)
        set_cell_text(row.cells[4], item.get("responsibility") or "\u4f01\u4e1a\u5b89\u5168\u7ba1\u7406\u90e8\u95e8", align=WD_ALIGN_PARAGRAPH.CENTER, font_size=9.5)

    if old_table is not None:
        old_table._tbl.addprevious(table._tbl)
        remove_element(old_table._tbl)


def find_last_paragraph_index(document, keyword):
    found = -1
    compact_keyword = keyword.replace(" ", "")
    for index, paragraph in enumerate(document.paragraphs):
        if compact_keyword in paragraph.text.replace(" ", ""):
            found = index
    return found


def find_paragraph_index_after(document, keyword, after_index):
    compact_keyword = keyword.replace(" ", "")
    for index in range(after_index + 1, len(document.paragraphs)):
        if compact_keyword in document.paragraphs[index].text.replace(" ", ""):
            return index
    return -1


def replace_reference_section(document, standards):
    start = find_last_paragraph_index(document, "\u68c0\u67e5\u4f9d\u636e")
    end = find_paragraph_index_after(document, "\uff08\u4e09\uff09\u73b0\u573a\u95ee\u9898\u9690\u60a3\u53ca\u6574\u6539\u5efa\u8bae\u63aa\u65bd\u6e05\u5355", start)
    if end < 0:
        end = find_paragraph_index_after(document, "\u73b0\u573a\u95ee\u9898\u9690\u60a3\u53ca\u6574\u6539\u5efa\u8bae\u63aa\u65bd\u6e05\u5355", start)
    if start < 0 or end < 0 or end <= start:
        return

    anchor = document.paragraphs[start]
    template_paragraph = next((p for p in document.paragraphs[start + 1:end] if clean(p.text)), anchor)
    for paragraph in list(document.paragraphs[start + 1:end]):
        remove_paragraph(paragraph)

    dynamic_standards = [normalize_reference_text(item) for item in standards or [] if normalize_reference_text(item)]
    if not dynamic_standards:
        dynamic_standards = ["未识别到明确排查依据，请结合知识库或现行标准人工复核补充；"]

    for index, standard in enumerate(dynamic_standards, start=1):
        anchor = insert_paragraph_after(anchor, f"{index}. {standard}", template_paragraph=template_paragraph)


def replace_opinion_section(document, payload):
    start = find_last_paragraph_index(document, LABELS["opinion_title"])
    if start < 0:
        return

    anchor = document.paragraphs[start]
    for paragraph in list(document.paragraphs[start + 1:]):
        remove_paragraph(paragraph)

    opinion = payload.get("comprehensiveOpinion") or {}
    directions = opinion.get("improvement_directions") or []
    general = clean(opinion.get("general_suggestions"))

    anchor = insert_paragraph_after(anchor, LABELS["opinion_sub_1"])
    for index, item in enumerate(directions, start=1):
        title = clean(item.get("title")) or f"\u6539\u8fdb\u65b9\u5411 {index}"
        content = clean(item.get("content"))
        anchor = insert_paragraph_after(anchor, f"{index}. {title}")
        if content:
            anchor = insert_paragraph_after(anchor, content, first_line_indent_cm=0.74)

    anchor = insert_paragraph_after(anchor, LABELS["opinion_sub_2"])
    if general:
        insert_paragraph_after(anchor, general, first_line_indent_cm=0.74)


def main():
    if len(sys.argv) != 4:
        print("usage: fillReportTemplate.py <template.docx> <output.docx> <payload.json>", file=sys.stderr)
        return 2

    template_path, output_path, payload_path = sys.argv[1:4]
    with open(payload_path, "r", encoding="utf-8") as handle:
        payload = json.load(handle)

    document = Document(template_path)
    enterprise = payload.get("enterprise", {})
    project_name = enterprise.get("project_name") or "\u9879\u76ee\u90e8"
    enterprise_name = enterprise.get("name") or LABELS["placeholder_company"]

    replace_all(document, {
        LABELS["placeholder_company"]: enterprise_name,
        LABELS["placeholder_project"]: project_name,
        LABELS["sample_project"]: project_name,
        LABELS["sample_project_reversed"]: project_name,
    })

    fill_cover_table(document, payload)
    replace_reference_section(document, payload.get("referenceStandards") or [])
    build_clean_hazard_table(document, payload)
    replace_opinion_section(document, payload)

    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    document.save(output_path)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
